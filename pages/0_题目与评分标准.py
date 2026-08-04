# -*- coding: utf-8 -*-
"""
pages/0_题目与评分标准.py — 作业模板管理（2026-07-20 常年化）
════════════════════════════════════════════════════════
换年级、换题目、换评分标准都在这里：
  粘贴评分表文字（或上传 docx）→ DeepSeek 解析成结构化标准 →
  老师逐项确认（判定方式可改）→ 保存为模板 → 主页建任务时选择。
原则：AI 解析结果必须过老师的眼才能保存——复核原则同样适用于规则本身。
"""

import json

import streamlit as st

import video_db as db
from video_rubric import (JUDGE_LABELS, dims, total_max, validate_rubric,
                          normalize_rubric, total_levels, member_mode)
from video_engine import parse_rubric_text

st.set_page_config(page_title="题目与评分标准", page_icon="📋", layout="wide")
if not st.session_state.get("authed"):
    st.warning("请先在主页输入平台口令")
    st.stop()

db.init_schema()
st.title("📋 题目与评分标准（作业模板）")
st.caption("每学期换题目/换标准：新建一个模板即可，历史批改记录不受影响。")

# ── 现有模板 ─────────────────────────────────────────────────
st.subheader("现有模板")
st.caption("换年级/换题目就新建一个；模板改错了可以直接改，不必删了重建——"
           "历史批改记录持有快照，改模板不影响已批改的成绩。")
rubrics = db.list_rubrics()
for r in rubrics:
    rb = r["rubric"]
    if isinstance(rb, str):
        rb = json.loads(rb)
    normalize_rubric(rb)
    dim_line = "、".join(f"{d['name']}{d['max']}分" for d in dims(rb))
    # 2026-08-04：标题行 + 操作按钮放在折叠条【外面】。
    # 原来删除按钮藏在 expander 里，老师不展开根本看不到（实测反馈）。
    hc1, hc2, hc3 = st.columns([7, 1.2, 1.2])
    with hc1:
        st.markdown(f"**#{r['id']} {r['name']}** ｜ 总分 {total_max(rb)} ｜ "
                    f"{dim_line}"
                    + ("　🧑‍🤝‍🧑小组逐人给分" if member_mode(rb) else ""))
    with hc2:
        edit_on = st.toggle("✏️ 编辑", key=f"edit_{r['id']}")
    with hc3:
        if len(rubrics) > 1:
            if st.session_state.get(f"confirm_del_{r['id']}"):
                if st.button("确认删除", key=f"do_del_{r['id']}",
                             type="primary"):
                    db.delete_rubric(r["id"])
                    st.session_state[f"confirm_del_{r['id']}"] = False
                    st.rerun()
            else:
                if st.button("🗑 删除", key=f"del_{r['id']}"):
                    st.session_state[f"confirm_del_{r['id']}"] = True
                    st.rerun()
        else:
            st.caption("（仅剩一个）")

    with st.expander("查看详情"):
        st.write(f"**题目要求默认文本**：{r['requirements'] or '（空）'}")
        if total_levels(rb):
            st.write("**总分等级表**：" + "　".join(
                f"{lv['grade']}{lv['label']} {lv['lo']}-{lv['hi']}"
                for lv in total_levels(rb)))
        if rb.get("reading_cap_total") is not None:
            st.write(f"**照读个人总分封顶**：{rb['reading_cap_total']} 分")
        for d in dims(rb):
            st.markdown(f"**{d['name']}（{d['max']}分）** — "
                        f"{JUDGE_LABELS.get(d.get('judge'), '?')}")
            for lv in d.get("levels", []):
                st.caption(f"{lv['grade']} {lv['label']}"
                           f"（{lv['lo']}–{lv['hi']}）：{lv['desc']}")
        pre = rb.get("precheck", {})
        if rb.get("stance"):
            st.write(f"**宽严指引**：{rb['stance']}")
        st.write(f"准入：时长≤{pre.get('max_duration_sec') or '不限'}秒 ｜ "
                 f"英文夹杂≤{round(pre.get('en_ratio_limit', 1.0)*100)}% ｜ "
                 f"出镜要求：{'是' if pre.get('face_required') else '否'} ｜ "
                 f"人工勾选：{'、'.join(pre.get('manual_items', [])) or '无'}")

    # ── 手工编辑（2026-08-04 新增）────────────────────────────
    # 起因：原来只有"粘贴原文→AI解析"一条路，而 DeepSeek 对整体量表
    # （每档同时描述四个面）会把档位小标题误认成维度——实测把 60 分四维
    # 表解析成"总分240、四个60分维度"。必须有一条不过 AI 的路。
    if edit_on:
        with st.form(f"form_{r['id']}"):
            e_name = st.text_input("模板名", value=r["name"])
            e_req = st.text_area("题目要求", value=r["requirements"] or "",
                                 height=120)
            e_stance = st.text_area("宽严指引（写给 AI 的评分基调）",
                                    value=rb.get("stance", ""), height=140)
            st.markdown("**各维度满分与判定方式**（改满分会自动重算总分）")
            judge_opts = list(JUDGE_LABELS.keys())
            for i, d in enumerate(dims(rb)):
                c1, c2, c3 = st.columns([3, 1.4, 2.6])
                with c1:
                    st.text_input("维度名", value=d["name"],
                                  key=f"dn_{r['id']}_{i}")
                with c2:
                    st.number_input("满分", min_value=1, max_value=200,
                                    value=int(d["max"]),
                                    key=f"dm_{r['id']}_{i}")
                with c3:
                    st.selectbox("判定方式", judge_opts,
                                 index=judge_opts.index(d.get("judge", "text")),
                                 format_func=lambda k: JUDGE_LABELS[k],
                                 key=f"dj_{r['id']}_{i}")
            st.markdown("**准入检查**")
            pc = rb.get("precheck", {})
            q1, q2, q3 = st.columns(3)
            with q1:
                e_dur = st.number_input("时长上限（秒，0=不限）", min_value=0,
                                        value=int(pc.get("max_duration_sec", 0)),
                                        key=f"pd_{r['id']}")
            with q2:
                e_en = st.number_input("英文夹杂上限（%）", min_value=0,
                                       max_value=100,
                                       value=int(round(
                                           pc.get("en_ratio_limit", 1.0) * 100)),
                                       key=f"pe_{r['id']}")
            with q3:
                e_face = st.checkbox("要求面容出镜",
                                     value=bool(pc.get("face_required")),
                                     key=f"pf_{r['id']}")
            e_manual = st.text_input(
                "人工勾选项（逗号分隔）",
                value="，".join(pc.get("manual_items", [])),
                key=f"pm_{r['id']}")
            st.caption("档位描述、总分等级表这类细节，需要改的话请用下面的"
                       "「高级：直接改 JSON」。")
            adv = st.text_area("高级：直接改 JSON（不懂就别动）",
                               value=json.dumps(rb, ensure_ascii=False,
                                                indent=1),
                               height=180, key=f"js_{r['id']}")
            saved = st.form_submit_button("💾 保存修改", type="primary")
        if saved:
            try:
                new_rb = json.loads(adv)
            except Exception as ex:
                st.error(f"JSON 格式有误，没有保存：{ex}")
            else:
                for i, d in enumerate(new_rb.get("dimensions", [])):
                    d["name"] = st.session_state.get(f"dn_{r['id']}_{i}", d["name"])
                    d["max"] = int(st.session_state.get(f"dm_{r['id']}_{i}", d["max"]))
                    d["judge"] = st.session_state.get(f"dj_{r['id']}_{i}", d.get("judge"))
                new_rb["stance"] = e_stance
                pre2 = new_rb.setdefault("precheck", {})
                pre2["max_duration_sec"] = int(e_dur)
                pre2["en_ratio_limit"] = e_en / 100.0
                pre2["face_required"] = bool(e_face)
                pre2["manual_items"] = [x.strip() for x in
                                        e_manual.replace(",", "，").split("，")
                                        if x.strip()]
                normalize_rubric(new_rb)
                new_rb["name"] = e_name.strip() or r["name"]
                errs = validate_rubric(new_rb)
                if errs:
                    for e in errs:
                        st.error(e)
                    st.info("以上问题不改掉不会保存——改坏的模板会让全班分数出错。")
                else:
                    db.update_rubric(r["id"], new_rb["name"], new_rb,
                                     e_req.strip())
                    st.success(f"模板「{new_rb['name']}」已更新。")
                    st.rerun()
    st.divider()

# ── 新建模板 ─────────────────────────────────────────────────
st.divider()
st.subheader("新建模板")
st.markdown("把新评分表**粘贴到下面**（从 Word 直接复制即可，表格格式乱没关系），"
            "或上传 docx 文件。解析需要 DeepSeek Key（主页侧栏填）。")

up_docx = st.file_uploader("上传评分表 docx（可选）", type=["docx"])
raw_default = ""
if up_docx is not None:
    try:
        from docx import Document
        doc = Document(up_docx)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        raw_default = "\n".join(parts)
        st.success(f"已读取 docx，共 {len(raw_default)} 字，"
                   f"内容已填入下方文本框，可先检查再解析。")
    except Exception as e:
        st.error(f"docx 读取失败：{e}，请改用粘贴方式。")

raw = st.text_area("评分表原文", value=raw_default, height=260,
                   key="rubric_raw",
                   placeholder="例：\n一、内容讲述（25分）\n优秀 21-25 步骤清晰完整…\n良好 16-20 …")
req_text = st.text_area("题目要求（会写进批改提示词，让 AI 知道作业要求什么）",
                        height=100,
                        placeholder="例：拍一个不超过3分钟的短视频，用华语介绍一本书…")

if st.button("🔍 AI 解析评分表", disabled=not raw.strip()):
    if not st.session_state.get("deepseek_key"):
        st.error("请先在主页侧栏填入 DeepSeek Key")
    else:
        with st.spinner("正在解析…"):
            try:
                rubric, errs = parse_rubric_text(
                    st.session_state.deepseek_key, raw)
                st.session_state.parsed_rubric = rubric
                st.session_state.parsed_errs = errs
            except Exception as e:
                st.error(f"解析失败：{e}")

# ── 解析结果确认 ─────────────────────────────────────────────
if st.session_state.get("parsed_rubric"):
    rubric = st.session_state.parsed_rubric
    st.divider()
    st.subheader("解析结果——请逐项核对后保存")
    for e in st.session_state.get("parsed_errs", []):
        st.error(f"体检未过：{e}（请修改上方原文重新解析）")

    name = st.text_input("模板名", value=rubric.get("name", ""))
    st.markdown(f"**总分 {total_max(rubric)} 分**，各维度如下。"
                f"请特别确认每个维度的**判定方式**（AI 预判，可改）：")
    judge_opts = list(JUDGE_LABELS.keys())
    for i, d in enumerate(dims(rubric)):
        c1, c2 = st.columns([3, 2])
        with c1:
            st.markdown(f"**{d['name']}（{d['max']}分）**"
                        + ("（重点）" if d.get("star") else ""))
            for lv in d.get("levels", []):
                st.caption(f"{lv['grade']} {lv['label']}"
                           f"（{lv['lo']}–{lv['hi']}）：{lv['desc']}")
        with c2:
            d["judge"] = st.selectbox(
                "判定方式", judge_opts,
                index=judge_opts.index(d.get("judge", "text")),
                format_func=lambda k: JUDGE_LABELS[k],
                key=f"judge_{i}")

    pre = rubric.setdefault("precheck", {})
    st.markdown("**准入检查配置**")
    p1, p2, p3 = st.columns(3)
    with p1:
        pre["max_duration_sec"] = st.number_input(
            "时长上限（秒，0=不限）", min_value=0, value=int(pre.get("max_duration_sec", 0)))
    with p2:
        pre["en_ratio_limit"] = st.number_input(
            "英文夹杂上限（%，100=不限）", min_value=0, max_value=100,
            value=int(round(pre.get("en_ratio_limit", 1.0) * 100))) / 100.0
    with p3:
        pre["face_required"] = st.checkbox(
            "要求面容出镜", value=bool(pre.get("face_required")))
    rubric["stance"] = st.text_area(
        "宽严指引（写给 AI 的评分基调，如：中一学生从宽、拍摄清楚即3分起步）",
        value=rubric.get("stance", ""), height=110)
    manual_str = st.text_input(
        "人工勾选项（老师复核时勾，逗号分隔）",
        value="，".join(pre.get("manual_items", [])))
    pre["manual_items"] = [s.strip() for s in
                           manual_str.replace(",", "，").split("，")
                           if s.strip()]

    if st.button("💾 保存为模板", type="primary", disabled=not name.strip()):
        normalize_rubric(rubric)
        rubric["name"] = name.strip()
        errs = validate_rubric(rubric)
        if errs:
            for e in errs:
                st.error(e)
        else:
            db.create_rubric(name.strip(), rubric, req_text.strip())
            st.session_state.parsed_rubric = None
            st.success(f"模板「{name.strip()}」已保存。到主页建任务时即可选择。")
            st.rerun()
